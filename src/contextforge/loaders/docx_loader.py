from __future__ import annotations

from pathlib import Path

from docx import Document

from .base import BaseLoader, LoadedDocument


class DocxLoader(BaseLoader):
    source_type = "docx"

    def load(self, path: Path, *, max_chars: int) -> LoadedDocument:
        doc = Document(str(path))
        paras = [p.text for p in doc.paragraphs if p.text]
        text = "\n".join(paras)
        if len(text) > max_chars:
            text = text[:max_chars]
        return LoadedDocument(
            source_path=str(path), source_type=self.source_type, title=None, text=text
        )
